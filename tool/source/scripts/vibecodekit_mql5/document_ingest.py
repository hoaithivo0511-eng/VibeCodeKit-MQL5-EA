"""Load EA requirements from PDF/DOCX/text and compile them to canonical EA-IR.

Document extraction is deliberately separate from semantic intake.  The loader
preserves PDF page boundaries with form-feed characters so every extracted
requirement can carry an auditable page-level source reference.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any


@dataclass(frozen=True)
class DocumentText:
    source: str
    format: str
    text: str
    page_count: int
    metadata: dict[str, Any]


def _clean_page(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def _load_pdf(path: Path) -> DocumentText:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:  # pragma: no cover - dependency contract
        raise RuntimeError("PDF intake requires pypdf>=5.0") from exc
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        try:
            unlocked = reader.decrypt("")
        except Exception as exc:  # noqa: BLE001
            raise ValueError("encrypted PDF cannot be read without a password") from exc
        if not unlocked:
            raise ValueError("encrypted PDF cannot be read without a password")
    pages: list[str] = []
    for page in reader.pages:
        pages.append(_clean_page(page.extract_text() or ""))
    meta = {str(k).lstrip("/"): str(v) for k, v in (reader.metadata or {}).items()}
    return DocumentText(
        source=str(path), format="pdf", text="\f".join(pages),
        page_count=len(pages), metadata=meta,
    )


def _table_text(table: Any) -> str:
    lines: list[str] = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def _load_docx(path: Path) -> DocumentText:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:  # pragma: no cover - core dependency
        raise RuntimeError("DOCX intake requires python-docx>=1.1") from exc
    doc = Document(str(path))
    chunks = [p.text for p in doc.paragraphs if p.text.strip()]
    chunks.extend(_table_text(t) for t in doc.tables)
    props = doc.core_properties
    meta = {
        "title": props.title or "",
        "subject": props.subject or "",
        "author": props.author or "",
        "version": props.version or "",
    }
    return DocumentText(
        source=str(path), format="docx", text="\n".join(chunks),
        page_count=1, metadata=meta,
    )


def _load_text(path: Path) -> DocumentText:
    text = path.read_text(encoding="utf-8", errors="replace")
    return DocumentText(
        source=str(path), format=path.suffix.lower().lstrip(".") or "text",
        text=text, page_count=max(1, text.count("\f") + 1), metadata={},
    )


def load_document(path: Path | str) -> DocumentText:
    path = Path(path)
    if not path.is_file():
        raise FileNotFoundError(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        doc = _load_pdf(path)
    elif suffix == ".docx":
        doc = _load_docx(path)
    elif suffix in {".txt", ".md", ".rst", ".csv", ".yaml", ".yml", ".json"}:
        doc = _load_text(path)
    else:
        raise ValueError(f"unsupported document format: {suffix or '<none>'}")
    if not doc.text.strip():
        raise ValueError(f"document contains no extractable text: {path}")
    return doc


def compile_document(path: Path | str, *, strict: bool = False):
    from .intake import parse_text

    doc = load_document(path)
    ir = parse_text(doc.text, source=doc.source, strict=strict)
    ir.metadata.update({
        "document_format": doc.format,
        "document_pages": doc.page_count,
        "document_metadata": doc.metadata,
    })
    return ir


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(prog="mql5-doc-intake-ir", description=__doc__.splitlines()[0])
    ap.add_argument("--file", required=True, type=Path)
    ap.add_argument("--out", type=Path)
    ap.add_argument("--strict", action="store_true")
    args = ap.parse_args(argv)
    try:
        ir = compile_document(args.file, strict=args.strict)
    except (OSError, ValueError, RuntimeError) as exc:
        print(f"mql5-doc-intake-ir: {exc}", file=sys.stderr)
        return 2
    payload = json.dumps(ir.to_dict(), ensure_ascii=False, indent=2) + "\n"
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(payload, encoding="utf-8")
    else:
        sys.stdout.write(payload)
    return 0 if ir.ready_for_planning else 2


if __name__ == "__main__":
    raise SystemExit(main())
