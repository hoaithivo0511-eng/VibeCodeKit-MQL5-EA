"""Generate senior-style EA review report DOCX."""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from .ea_senior_review import review_project

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.section import WD_ORIENT
except Exception:  # pragma: no cover
    Document = None
    Inches = None
    Pt = None
    WD_ORIENT = None


def write_review_docx(report: dict[str, Any], out: str | Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not available")
    out = Path(out)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    doc.add_heading("Senior EA Code Review", 0)
    doc.add_paragraph("Static strategy/risk/execution review for an existing MQL5 EA codebase.")
    doc.add_paragraph(f"Project: {report['project']}")
    doc.add_paragraph(f"Score: {report['score']}/100")
    doc.add_paragraph(f"Readiness: {report['readiness']}")

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(report.get("senior_summary", ""))

    doc.add_heading("2. Strategy classification", level=1)
    strategy = report.get("strategy", {})
    doc.add_paragraph(f"Detected family: {strategy.get('family')}")
    doc.add_paragraph("Detected signals:")
    for s in strategy.get("signals", []):
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("3. Architecture/features observed", level=1)
    features = report.get("analysis_summary", {}).get("features", {})
    for k, v in features.items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    doc.add_heading("4. Findings by severity", level=1)
    issues = report.get("issues", [])
    if not issues:
        doc.add_paragraph("No static issue found. Real compile/backtest validation is still required.")
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Severity", "Category", "Finding", "Evidence", "Recommendation"]
        for idx, h in enumerate(headers):
            table.rows[0].cells[idx].text = h
        for issue in issues:
            row = table.add_row().cells
            row[0].text = issue.get("severity", "")
            row[1].text = issue.get("category", "")
            row[2].text = issue.get("title", "")
            row[3].text = issue.get("evidence", "")
            row[4].text = issue.get("recommendation", "")
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

    doc.add_heading("5. Senior developer recommendations", level=1)
    crit = [i for i in issues if i.get("severity") == "critical"]
    err = [i for i in issues if i.get("severity") == "error"]
    warn = [i for i in issues if i.get("severity") == "warn"]
    if crit:
        doc.add_paragraph("Fix these before any release attempt:", style="List Bullet")
        for i in crit:
            doc.add_paragraph(f"{i['title']}: {i['recommendation']}", style="List Number")
    if err:
        doc.add_paragraph("Fix these before serious forward test:", style="List Bullet")
        for i in err:
            doc.add_paragraph(f"{i['title']}: {i['recommendation']}", style="List Number")
    if warn:
        doc.add_paragraph("Recommended improvements:", style="List Bullet")
        for i in warn:
            doc.add_paragraph(f"{i['title']}: {i['recommendation']}", style="List Number")

    doc.add_heading("6. Release readiness note", level=1)
    doc.add_paragraph("This report is static review only. It does not replace MetaEditor compile, MT5 Strategy Tester, multi-broker validation, walk-forward testing, or live forward testing.")

    doc.save(out)


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Generate senior EA review DOCX report.")
    ap.add_argument("--project", required=True)
    ap.add_argument("--profile", default="auto")
    ap.add_argument("--out", required=True)
    ap.add_argument("--json-out")
    args = ap.parse_args(argv)

    report = review_project(args.project, args.profile)
    write_review_docx(report, args.out)
    if args.json_out:
        p = Path(args.json_out)
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({"ok": True, "docx": args.out, "json": args.json_out, "score": report["score"], "readiness": report["readiness"]}, indent=2, ensure_ascii=False))
    return 0 if report["readiness"] != "release-blocked" else 2


if __name__ == "__main__":
    raise SystemExit(main())
