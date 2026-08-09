from pathlib import Path

import pytest
from docx import Document

from vibecodekit_mql5.document_ingest import compile_document, load_document


def test_text_document_preserves_page_boundaries(tmp_path: Path):
    path = tmp_path / "manual.txt"
    path.write_text(
        "EA named AtlasDCA account hedging EURUSD H1 DCA Step Multiplier\f"
        "standard hedge and max spread 2",
        encoding="utf-8",
    )
    doc = load_document(path)
    assert doc.page_count == 2
    ir = compile_document(path, strict=True)
    assert ir.runtime["account_model"] == "hedging"
    assert "strategy.dca.step_multiplier" in ir.strategy["features"]
    refs = [ref for req in ir.requirements for ref in req.source_refs]
    assert {ref.page for ref in refs} >= {1, 2}


def test_docx_tables_are_part_of_intake(tmp_path: Path):
    path = tmp_path / "strategy.docx"
    doc = Document()
    doc.add_paragraph("EA named TableEA account netting EURUSD H1 breakout ATR break")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Parameter"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "max spread"
    table.cell(1, 1).text = "2.5"
    doc.save(path)
    ir = compile_document(path, strict=True)
    assert "strategy.entry.breakout" in ir.strategy["features"]
    assert ir.risk["max_spread_pips"] == 2.5
    assert ir.metadata["document_format"] == "docx"


def test_pdf_loader_is_generic_not_ccbsn_specific(tmp_path: Path):
    reportlab = pytest.importorskip("reportlab.pdfgen.canvas")
    path = tmp_path / "pulse.pdf"
    canvas = reportlab.Canvas(str(path))
    canvas.drawString(40, 800, "EA named PulseBreak account netting XAUUSD M15 breakout ATR break")
    canvas.drawString(40, 780, "base lot 0.01 max lot 1 max spread 2.5 max positions 4")
    canvas.save()
    ir = compile_document(path, strict=True)
    assert ir.identity["name"] == "PulseBreak"
    assert ir.runtime["account_model"] == "netting"
    assert "strategy.entry.breakout" in ir.strategy["features"]
    assert "strategy.dca.enabled" not in ir.strategy["features"]
    assert ir.metadata["document_pages"] == 1


def test_unsupported_binary_format_is_rejected(tmp_path: Path):
    path = tmp_path / "manual.bin"
    path.write_bytes(b"EA named Nope")
    with pytest.raises(ValueError, match="unsupported document format"):
        load_document(path)
