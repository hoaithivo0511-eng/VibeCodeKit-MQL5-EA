"""Single-command deep-review orchestrator (``mql5-ea-deep-review``).

UX principle (v2.4): the tool surface grows by exactly ONE command. A user
(or an agent prompt like \"scan / deep-review / audit this EA code-base\")
runs this one entry point and the orchestrator executes the whole pipeline,
then emits ONE unified report.

Pipeline stages
---------------
    Stage 0  parse + symbol graph        (mq5_symbols)
    Stage 1  static scan / signals       (scan_ea)
    Stage 2  anti-pattern lint           (lint + lint_best_practice)
    Stage 3  structure audit             (structure_audit)   \
    Stage 4  dead-code / dead-logic       (deadcode)          | via senior
    Stage 5  senior review + risk/release (ea_senior_review)  /  review
    Stage 6  modernization advisor        (modernize)
    Stage 7  grounded packet preparation  (line_review; no LLM verdict claimed)

Stages 3-5 are produced by :func:`ea_senior_review.review_project`, which
already folds the structure-audit + dead-code layers into its
``code_quality`` category, so the orchestrator never double-counts them.

Flags
-----
    (default)        Stages 0-6 + Stage-7 packet, Markdown/JSON/DOCX
    --fast           skip the grounded line-review packet build (Stage 7)
    --json-only      do not write the DOCX report
    --no-docx        alias of --json-only for the document artifact
    --profile NAME   forward a risk profile to the senior review
    --out DIR        output directory for the unified report artifacts
"""
from __future__ import annotations

import argparse
import json
import sys
import tempfile
from pathlib import Path
from typing import Any

from . import _agent_io
from .ea_doc_analyzer import read_mql_files
from .ea_senior_review import review_project
from .line_review import run_line_review
from .lint import lint_source
from .modernize import analyze_modernization
from .scan_ea import analyze_source

TOOL = "mql5-ea-deep-review"

_SEV_FROM_LINT = {"ERROR": "error", "WARN": "warn", "INFO": "info"}
_SEV_RANK = {"critical": 0, "error": 1, "warn": 2, "info": 3}

_STAGE_NAMES = {
    0: "parse and symbol graph",
    1: "static scan and signals",
    2: "anti-pattern lint",
    3: "structure and complexity",
    4: "dead-code and dead-logic",
    5: "senior risk, execution, state and release review",
    6: "modernization advisor",
    7: "grounded line-review packet preparation",
}


def _stage(stage_id: int, status: str, detail: str) -> dict[str, Any]:
    return {
        "id": stage_id,
        "name": _STAGE_NAMES[stage_id],
        "status": status,
        "detail": detail,
    }


def _resolve_files(target: Path) -> tuple[Path, dict[str, str]]:
    """Return (project_root, {relpath: text}) for a file or directory."""
    if target.is_file():
        root = target.parent
        return root, {target.name: _read(target)}
    return target, read_mql_files(target)


def _read(path: Path) -> str:
    from .mq5_io import read_mq5_text
    return read_mq5_text(path, errors="replace")


def _lint_issues(files: dict[str, str]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for rel, text in files.items():
        for f in lint_source(rel, text):
            out.append({
                "severity": _SEV_FROM_LINT.get(f.severity, "warn"),
                "category": "anti_pattern",
                "title": f"{f.code}: {f.message}",
                "evidence": f"[{rel}] line {f.line}:{f.col} {f.message}",
                "recommendation": "See kit anti-pattern table for remediation.",
                "line": f.line,
                "file": rel,
            })
    return out


def _modernize_issues(files: dict[str, str]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for rel, text in files.items():
        for it in analyze_modernization(text)["issues"]:
            it = dict(it)
            it["file"] = rel
            it["evidence"] = f"[{rel}] " + it["evidence"]
            out.append(it)
    return out


def _line_review_section(files: dict[str, str], static_by_file: dict[str, list],
                         fast: bool) -> dict[str, Any]:
    if fast:
        return {"mode": "skipped", "reason": "--fast", "packets": 0}
    total_packets = 0
    packs: list[str] = []
    for rel, text in files.items():
        res = run_line_review(text, file=rel,
                              static_findings=static_by_file.get(rel, []))
        total_packets += len(res.get("packets", []))
        if res.get("paste_pack"):
            packs.append(res["paste_pack"])
    return {"mode": "pack", "packets": total_packets,
            "paste_pack": "\n\n".join(packs)}


def run_deep_review(target: str | Path, *, profile: str = "auto",
                    fast: bool = False) -> dict[str, Any]:
    target = Path(target)
    root, files = _resolve_files(target)
    if not files:
        return {"ok": False, "error": f"no .mq5/.mqh sources under {target}"}
    stages: list[dict[str, Any]] = []

    # When the target is a single file, isolate it so the senior review
    # (which recursively reads every .mq5/.mqh under its project root) scans
    # ONLY the requested EA, never sibling files in the same folder.
    isolation = None
    if target.is_file():
        isolation = tempfile.TemporaryDirectory(prefix="mql5_deep_review_")
        for rel, text in files.items():
            dst = Path(isolation.name) / rel
            dst.parent.mkdir(parents=True, exist_ok=True)
            dst.write_text(text, encoding="utf-8")
        senior_root = Path(isolation.name)
    else:
        senior_root = root

    # Stage 1: scan / signal summary (informational)
    scan: dict[str, Any] = {}
    for rel, text in files.items():
        rep = analyze_source(text, source=rel)
        scan[rel] = rep.__dict__ if hasattr(rep, "__dict__") else {}
    stages.append(_stage(1, "EXECUTED", f"{len(files)} source file(s) scanned"))

    # Stage 3-5: senior review (includes code_quality = structure + deadcode)
    try:
        senior = review_project(senior_root, profile=profile)
    finally:
        if isolation is not None:
            isolation.cleanup()
    issues: list[dict[str, Any]] = list(senior.get("issues", []))
    stages.extend(
        [
            _stage(3, "EXECUTED", "structure metrics produced by senior review"),
            _stage(4, "EXECUTED", "dead-code analysis produced by senior review"),
            _stage(5, "EXECUTED", "senior review completed"),
        ]
    )

    # Stage 2: anti-pattern lint
    issues.extend(_lint_issues(files))
    stages.append(_stage(2, "EXECUTED", "anti-pattern detectors completed"))
    # Stage 6: modernization advisor
    issues.extend(_modernize_issues(files))
    stages.append(_stage(6, "EXECUTED", "modernization analysis completed"))

    # group static code_quality findings per file for line-review grounding
    static_by_file: dict[str, list] = {}
    for it in issues:
        if it.get("category") == "code_quality" and it.get("file"):
            static_by_file.setdefault(it["file"], []).append(it)

    # Stage 7 prepares grounded review packets; it does not fabricate an LLM verdict.
    line_review = _line_review_section(files, static_by_file, fast)
    if line_review.get("mode") == "skipped":
        stages.append(_stage(7, "SKIPPED", str(line_review.get("reason", "not run"))))
    else:
        stages.append(
            _stage(
                7,
                "EXECUTED",
                f"{line_review.get('packets', 0)} grounded packet(s) prepared; no LLM verdict claimed",
            )
        )

    # v2.5 (#6): structured graphs (call / input-usage / order-lifecycle /
    # risk-invariant). Additive evidence section; does not alter the senior
    # score/issue counts so existing contracts stay stable.
    from .mq5_graphs import analyze_graphs
    graphs = analyze_graphs(files)
    stages.append(_stage(0, "EXECUTED", "symbol and lifecycle graphs generated"))
    stages.sort(key=lambda item: item["id"])
    checked_categories = [
        f"Stage {item['id']}: {item['name']}"
        for item in stages
        if item["status"] == "EXECUTED"
    ]

    counts = {s: sum(1 for i in issues if i.get("severity") == s)
              for s in ("critical", "error", "warn", "info")}
    by_category: dict[str, int] = {}
    for i in issues:
        by_category[i.get("category", "other")] = by_category.get(
            i.get("category", "other"), 0) + 1

    issues_sorted = sorted(
        issues,
        key=lambda x: (_SEV_RANK.get(x.get("severity"), 9),
                       x.get("category", ""), x.get("line") or 0),
    )

    return {
        "schema_version": "2.5",
        "artifact_type": "ea_deep_review",
        "tool": TOOL,
        "project": str(root),
        "files_scanned": sorted(files.keys()),
        "strategy": senior.get("strategy", {}),
        "score": senior.get("score"),
        "readiness": senior.get("readiness"),
        "issue_counts": counts,
        "issue_categories": by_category,
        "issues": issues_sorted,
        "code_quality": senior.get("analysis_summary", {}).get("code_quality", {}),
        "line_review": line_review,
        "scan": scan,
        "graphs": graphs,
        "stages": stages,
        "checked_categories": checked_categories,
    }


def render_markdown(report: dict[str, Any]) -> str:
    lines: list[str] = []
    a = lines.append
    a(f"# Deep Review — {Path(report['project']).name}")
    a("")
    a(f"- Readiness: **{report.get('readiness')}**  |  Score: "
      f"**{report.get('score')}/100**")
    strat = report.get("strategy", {})
    a(f"- Strategy: {strat.get('family', 'n/a')} "
      f"({', '.join(strat.get('signals', []))})")
    a(f"- Strategy detection: {strat.get('detection_source', 'n/a')}")
    a(f"- Files scanned: {len(report.get('files_scanned', []))}")
    cq = report.get("code_quality", {})
    if cq:
        a(f"- Code metrics: {cq.get('function_count', 0)} functions, "
          f"max complexity {cq.get('max_complexity', 0)}, "
          f"{cq.get('dead_findings', 0)} dead-code findings")
    a("")
    a("## Stage execution")
    a("| Stage | Status | Detail |")
    a("|---|---|---|")
    for stage in report.get("stages", []):
        a(
            f"| {stage.get('id')}: {stage.get('name')} | "
            f"{stage.get('status')} | {stage.get('detail')} |"
        )
    a("")
    counts = report.get("issue_counts", {})
    a("## Issue summary")
    a(f"- Critical: {counts.get('critical', 0)}")
    a(f"- Error: {counts.get('error', 0)}")
    a(f"- Warn: {counts.get('warn', 0)}")
    a(f"- Info: {counts.get('info', 0)}")
    a("")
    a("### By category")
    for cat, n in sorted(report.get("issue_categories", {}).items()):
        a(f"- {cat}: {n}")
    a("")
    a("## Findings")
    cur = None
    for it in report.get("issues", []):
        sev = it.get("severity", "info")
        if sev != cur:
            cur = sev
            a(f"\n### {sev.upper()}")
        loc = f" (line {it['line']})" if it.get("line") else ""
        a(f"- **{it.get('title')}**{loc} — {it.get('evidence')}")
        if it.get("recommendation"):
            a(f"  - Fix: {it['recommendation']}")
    lr = report.get("line_review", {})
    a("")
    a("## Grounded line review")
    if lr.get("mode") == "skipped":
        a("- Skipped (--fast).")
    else:
        a(f"- {lr.get('packets', 0)} function packets prepared for grounded "
          "LLM review (paste-pack available in the JSON artifact).")
    return "\n".join(lines)


def _write_docx(md: str, path: Path) -> bool:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return False
    doc = Document()
    for line in md.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(str(path))
    return True


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        prog=TOOL,
        description="One-command deep audit/review of an MQL5 EA code-base.")
    ap.add_argument("target", type=Path,
                    help="Path to a .mq5 file or an EA project directory.")
    ap.add_argument("--profile", default="auto",
                    help="Risk profile forwarded to the senior review.")
    ap.add_argument("--fast", action="store_true",
                    help="Skip the Stage-7 line-review packet build.")
    ap.add_argument("--out", type=Path, default=None,
                    help="Directory for unified report artifacts.")
    docx_grp = ap.add_mutually_exclusive_group()
    docx_grp.add_argument("--no-docx", action="store_true",
                          help="Do not write the DOCX report.")
    docx_grp.add_argument("--json-only", action="store_true",
                          help="Alias of --no-docx.")
    _agent_io.add_json_flag(ap)
    args = ap.parse_args(argv)

    if not args.target.exists():
        env = _agent_io.Envelope(tool=TOOL, ok=False, exit_code=2,
                                 summary=f"target not found: {args.target}",
                                 data={"target": str(args.target)})
        _agent_io.maybe_emit(args, env)
        if not getattr(args, "emit_json", False):
            sys.stderr.write(f"error: target not found: {args.target}\n")
        return 2

    report = run_deep_review(args.target, profile=args.profile, fast=args.fast)
    if not report.get("issue_counts") and report.get("error"):
        env = _agent_io.Envelope(tool=TOOL, ok=False, exit_code=2,
                                 summary=report["error"], data=report)
        _agent_io.maybe_emit(args, env)
        if not getattr(args, "emit_json", False):
            sys.stderr.write("error: " + report["error"] + "\n")
        return 2

    md = render_markdown(report)
    out_dir = args.out or (Path(report["project"]) / "_deep_review")
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "deep-review.md").write_text(md, encoding="utf-8")
    (out_dir / "deep-review.json").write_text(
        json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    docx_written = False
    if not (args.no_docx or args.json_only):
        docx_written = _write_docx(md, out_dir / "deep-review.docx")

    blocked = report.get("readiness") == "release-blocked"
    summary = (f"Deep review complete: {report['issue_counts']} "
               f"(readiness={report.get('readiness')}).")
    env = _agent_io.Envelope(
        tool=TOOL, ok=not blocked, exit_code=1 if blocked else 0,
        summary=summary,
        data={
            "readiness": report.get("readiness"),
            "score": report.get("score"),
            "issue_counts": report["issue_counts"],
            "issue_categories": report["issue_categories"],
            "out_dir": str(out_dir),
            "docx": docx_written,
            "line_review_packets": report.get("line_review", {}).get("packets", 0),
            "stages": report.get("stages", []),
        },
        evidence=report.get("files_scanned", []),
    )
    if getattr(args, "emit_json", False):
        _agent_io.emit(env)
    else:
        print(md)
        print(f"\n[artifacts written to {out_dir}]")
    return env.exit_code


if __name__ == "__main__":
    raise SystemExit(main())
